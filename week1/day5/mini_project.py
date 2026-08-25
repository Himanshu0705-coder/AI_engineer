import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key kaha hai bhai")

client=Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"


job_description="""
Description
# Data Analyst – Fresher / Entry Level

## Job Overview

We are looking for a motivated and analytical **Data Analyst** to join our team. The ideal candidate will be responsible for collecting, cleaning, analyzing, and interpreting data to support business decisions. This role is suitable for a fresher or recent graduate with a strong foundation in **SQL, Python, Excel, data visualization, and statistical analysis**.

The candidate should be comfortable working with large datasets, identifying trends and patterns, creating dashboards, and communicating actionable insights to business stakeholders.

## Key Responsibilities

* Collect, clean, transform, and validate data from multiple sources.
* Analyze structured and unstructured datasets to identify trends, patterns, and business insights.
* Write and optimize **SQL queries** for data extraction and analysis.
* Perform **Exploratory Data Analysis (EDA)** using Python and relevant libraries.
* Develop interactive dashboards and reports using **Power BI or Tableau**.
* Use **Excel** for data analysis, reporting, pivot tables, lookups, and data visualization.
* Create and monitor **KPIs and business metrics**.
* Perform statistical analysis and identify relationships, trends, and anomalies in datasets.
* Automate repetitive data-processing and reporting tasks using Python where appropriate.
* Work with business stakeholders to understand analytical requirements and translate them into data-driven solutions.
* Present analytical findings and recommendations to technical and non-technical stakeholders.
* Ensure data accuracy, consistency, and quality throughout the analysis process.
* Document data analysis methodologies, reports, dashboards, and insights.

## Required Qualifications

* Bachelor's degree in **Computer Science, Information Technology, Data Science, Statistics, Mathematics, Engineering, Business Analytics**, or a related field.
* Freshers and recent graduates are encouraged to apply.
* Strong understanding of **SQL and relational databases**.
* Good knowledge of **Python** for data analysis.
* Working knowledge of **Pandas, NumPy, Matplotlib, and Seaborn**.
* Proficiency in **Microsoft Excel**, including Pivot Tables, VLOOKUP/XLOOKUP, INDEX-MATCH, and basic data analysis functions.
* Knowledge of **Power BI or Tableau**.
* Understanding of data cleaning, data transformation, EDA, and data visualization.
* Basic understanding of **statistics and analytical concepts**.
* Strong problem-solving and analytical thinking skills.
* Good verbal and written communication skills.

## Preferred Skills

* Experience with **MySQL, PostgreSQL, or SQL Server**.
* Knowledge of **Power Query and DAX**.
* Understanding of SQL **JOINs, CTEs, subqueries, CASE statements, aggregate functions, and window functions**.
* Knowledge of KPI development and business reporting.
* Familiarity with Git/GitHub.
* Understanding of basic ETL/data-processing concepts.
* Experience working with large datasets.
* Knowledge of cloud platforms such as **AWS, Azure, or Google Cloud** is a plus.

## Technical Skills

**Programming & Querying:**
Python, SQL

**Data Analysis:**
Pandas, NumPy, Data Cleaning, Data Transformation, EDA, Statistical Analysis, Trend Analysis

**Visualization & BI:**
Power BI, Tableau, Power Query, DAX, Data Visualization, Dashboard Development

**Excel:**
Pivot Tables, Pivot Charts, VLOOKUP, XLOOKUP, INDEX-MATCH, Conditional Formatting

**Databases:**
MySQL, PostgreSQL, SQL Server

**Tools:**
Git, GitHub, Jupyter Notebook, VS Code

## Key Competencies

* Analytical Thinking
* Problem Solving
* Attention to Detail
* Data-driven Decision Making
* Business Acumen
* Communication Skills
* Stakeholder Management
* Time Management
* Team Collaboration

## What We Offer

* Opportunity to work on real-world business datasets.
* Exposure to data analytics and business intelligence projects.
* Mentorship from experienced analysts.
* Opportunities to develop technical and business skills.
* Collaborative and learning-oriented work environment.

"""
class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""
Analyze the following job description:

{job_description}
"""
message_system={
    "role" : "system",
    "content" : system_prompt
}
message_user={
    "role" : "user",
    "content" : user_prompt
}
response_format={
    "type" : "json_object"
}


messages=[message_system, message_user]

response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)


answer=response.choices[0].message.content

raw_json=answer
# print(raw_json)



import json
job_data=json.loads(raw_json)

job = JobD(**job_data)

print(job.minimum_experience)
print(job.education_requirements)



#parse real
class MatchResult(BaseModel):
    score: float
    details: dict
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


resume_schema = Resume.model_json_schema()
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None



# lets do it now
resume_folder = Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():
    #C:\Users\Pratyush\padho_with_pratyush\week1\day5\resumes\abhay resume new - Abhay Singh.pdf
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume=parse_resume(resume_text) # llm call1
    time.sleep(5)
    result = final_score(job, parsed_resume) #llm caLL2
    #score and details
    #acount chtgpt
    # request bhejna shhur krega millions
    #chattgot server jam ho jayega
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })
all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)
top_2 = all_results[:2]
worst_2 = all_results[-2:]


print("TOP 2 CANDIDATES")
for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])